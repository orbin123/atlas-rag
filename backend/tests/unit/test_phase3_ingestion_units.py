from __future__ import annotations

import asyncio
import zipfile
from pathlib import Path

import pytest
from docx import Document
from phase3_support import (
    CountingEncoderFactory,
    DeterministicEncoder,
    phase3_settings,
    write_text_pdf,
)

from app.services.embedding import EmbeddingService
from app.services.parsers import ParsedPage, ParserError, parse_document
from app.services.preprocessing import clean_pages
from app.storage.files import (
    StagedUpload,
    UploadFileError,
    validate_filename,
    validate_staged_content,
)


def test_real_txt_docx_and_pdf_parsers_preserve_truthful_pages(tmp_path: Path) -> None:
    text_path = tmp_path / "latin.txt"
    text_path.write_bytes("Caf\xe9 Atlas\r\nsecond line".encode("latin-1"))
    assert parse_document(text_path, "txt")[0].raw_text == "Café Atlas\nsecond line"

    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Atlas paragraph")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "left"
    table.cell(0, 1).text = "right"
    document.save(str(docx_path))
    docx_pages = parse_document(docx_path, "docx")
    assert len(docx_pages) == 1
    assert "Atlas paragraph" in docx_pages[0].raw_text
    assert "left\tright" in docx_pages[0].raw_text

    pdf_path = tmp_path / "sample.pdf"
    write_text_pdf(pdf_path, "Atlas PDF evidence")
    pdf_pages = parse_document(pdf_path, "pdf")
    assert len(pdf_pages) == 1
    assert pdf_pages[0].page_number == 1
    assert "Atlas PDF evidence" in pdf_pages[0].raw_text


def test_parser_rejects_empty_or_binary_text(tmp_path: Path) -> None:
    path = tmp_path / "empty.txt"
    path.write_text("  \n")
    with pytest.raises(ParserError, match="no extractable text"):
        parse_document(path, "txt")
    path.write_bytes(b"text\x00binary")
    with pytest.raises(ParserError, match="binary null"):
        parse_document(path, "txt")


def test_cleaning_removes_repeated_margins_and_normalizes_deterministically() -> None:
    pages = [
        ParsedPage(1, "ATLAS REPORT\nFirst hy-\nphenated fact.\nPage footer"),
        ParsedPage(2, "ATLAS REPORT\nSecond  fact.\nPage footer"),
    ]
    first = clean_pages(pages, remove_repeated_margins=True)
    second = clean_pages(pages, remove_repeated_margins=True)

    assert first == second
    assert first[0].repeated_lines_removed == ("ATLAS REPORT", "Page footer")
    assert "hyphenated" in first[0].cleaned_text
    assert "ATLAS REPORT" not in first[0].cleaned_text


def test_filename_mime_signature_and_docx_archive_safety(tmp_path: Path) -> None:
    with pytest.raises(UploadFileError) as traversal:
        validate_filename("../unsafe.txt")
    assert traversal.value.code == "VALIDATION_ERROR"
    with pytest.raises(UploadFileError) as unsupported:
        validate_filename("sample.exe")
    assert unsupported.value.status_code == 415

    settings = phase3_settings(tmp_path)
    fake_pdf = tmp_path / "fake.pdf"
    fake_pdf.write_bytes(b"not a pdf")
    staged_pdf = StagedUpload("fake.pdf", "pdf", "application/pdf", fake_pdf, 9, "a" * 64)
    with pytest.raises(UploadFileError, match="signature"):
        validate_staged_content(staged_pdf, settings)

    unsafe_docx = tmp_path / "unsafe.docx"
    with zipfile.ZipFile(unsafe_docx, "w") as archive:
        archive.writestr("[Content_Types].xml", "types")
        archive.writestr("word/document.xml", "document")
        archive.writestr("../escape.txt", "escape")
    staged_docx = StagedUpload(
        "unsafe.docx",
        "docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        unsafe_docx,
        unsafe_docx.stat().st_size,
        "b" * 64,
    )
    with pytest.raises(UploadFileError, match="unsafe archive"):
        validate_staged_content(staged_docx, settings)


def test_embedding_service_loads_one_shared_encoder(tmp_path: Path) -> None:
    settings = phase3_settings(tmp_path)
    encoder = DeterministicEncoder(settings.embedding_dimension)
    factory = CountingEncoderFactory(encoder)

    async def exercise() -> None:
        service = EmbeddingService(settings, factory)
        first, second = await asyncio.gather(service.load(), service.load())
        assert first is second
        vectors = await service.encode(["one", "two"], batch_size=1)
        assert vectors.shape == (2, 4)
        assert service.ready

    asyncio.run(exercise())
    assert factory.calls == 1
    assert encoder.encode_calls == 1
