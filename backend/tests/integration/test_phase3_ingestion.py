from __future__ import annotations

import hashlib
import time
from pathlib import Path
from typing import cast
from unittest.mock import patch

from docx import Document as DocxDocument
from fastapi import FastAPI
from fastapi.testclient import TestClient
from filelock import FileLock
from phase3_support import (
    CountingEncoderFactory,
    DeterministicEncoder,
    phase3_settings,
    write_text_pdf,
)
from sqlalchemy import func, select

from app.core.config import Settings
from app.db.migrations import upgrade_database
from app.db.models import Chunk, Document, DocumentPage, IndexState, IngestionJob
from app.db.session import create_database_engine, create_session_factory
from app.main import create_app


def _wait_for_job(client: TestClient, job_id: str, *, timeout: float = 8) -> dict[str, object]:
    deadline = time.monotonic() + timeout
    last: dict[str, object] = {}
    while time.monotonic() < deadline:
        response = client.get(f"/api/v1/ingestion-jobs/{job_id}")
        assert response.status_code == 200
        last = response.json()
        if last["status"] in {"succeeded", "failed", "cancelled"}:
            return last
        time.sleep(0.02)
    raise AssertionError(f"Job did not finish before timeout: {last}")


def _app(client: TestClient) -> FastAPI:
    return cast(FastAPI, client.app)


def _upload(
    client: TestClient,
    name: str,
    content: bytes,
    mime: str,
    **data: str,
) -> tuple[str, str]:
    response = client.post(
        "/api/v1/documents/upload",
        files={"file": (name, content, mime)},
        data=data,
    )
    assert response.status_code == 202, response.text
    payload = response.json()
    assert payload["statusUrl"] == f"/api/v1/ingestion-jobs/{payload['jobId']}"
    return str(payload["jobId"]), str(payload["documentId"])


def test_upload_all_formats_incrementally_indexes_and_survives_restart(tmp_path: Path) -> None:
    settings = phase3_settings(tmp_path)
    encoder = DeterministicEncoder(settings.embedding_dimension)
    factory = CountingEncoderFactory(encoder)
    docx_path = tmp_path / "evidence.docx"
    docx = DocxDocument()
    docx.add_paragraph("DOCX evidence explains the durable Atlas ingestion path.")
    docx.save(str(docx_path))
    pdf_path = tmp_path / "evidence.pdf"
    write_text_pdf(pdf_path, "PDF evidence is indexed with a real page number.")

    app = create_app(settings, encoder_factory=factory, start_worker=True)
    with TestClient(app) as client:
        initial_ready = client.get("/health/ready").json()
        assert initial_ready["checks"]["embedding"]["ready"] is True
        assert initial_ready["checks"]["worker"]["ready"] is True
        uploads = [
            (
                "evidence.txt",
                b"TXT evidence describes incremental indexing and durable restart behavior.",
                "text/plain",
                {"domain": "Testing", "title": "TXT evidence"},
            ),
            (
                "evidence.docx",
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                {"author": "Atlas"},
            ),
            ("evidence.pdf", pdf_path.read_bytes(), "application/pdf", {}),
        ]
        document_ids: list[str] = []
        for name, content, mime, metadata in uploads:
            job_id, document_id = _upload(client, name, content, mime, **metadata)
            terminal = _wait_for_job(client, job_id)
            assert terminal["status"] == "succeeded", terminal
            assert terminal["stage"] == "finalizing"
            assert terminal["progressPercent"] == 100
            result = terminal["result"]
            assert isinstance(result, dict)
            assert int(result["chunkCount"]) >= 1
            document_ids.append(document_id)

        assert factory.calls == 1
        assert client.get("/health/ready").status_code == 200
        stats = client.get("/api/v1/documents/stats").json()
        assert stats["totalDocuments"] == 3
        assert stats["indexedDocuments"] == 3
        assert stats["indexHealth"]["vectorCount"] == stats["totalChunks"]
        for document_id in document_ids:
            detail = client.get(f"/api/v1/documents/{document_id}").json()
            assert detail["status"] == "indexed"
            assert detail["sourceKind"] == "upload"
            assert detail["relativeSourcePath"] is None
            assert client.get(f"/api/v1/documents/{document_id}/pages").json()["total"] >= 1
            assert client.get(f"/api/v1/documents/{document_id}/chunks").json()["total"] >= 1

        duplicate = client.post(
            "/api/v1/documents/upload",
            files={"file": (uploads[0][0], uploads[0][1], uploads[0][2])},
        )
        assert duplicate.status_code == 409
        assert duplicate.json()["error"]["code"] == "DUPLICATE_DOCUMENT"
        assert duplicate.json()["error"]["details"]["documentId"] == document_ids[0]

        session_factory = _app(client).state.session_factory
        with session_factory() as session:
            row = session.scalar(select(Chunk).where(Chunk.document_id == document_ids[0]).limit(1))
            assert row is not None
            query_text = row.cleaned_text
            expected_vector_id = row.vector_id
            state = session.get_one(IndexState, 1)
            version_before_restart = state.index_version

    restarted = create_app(settings, encoder_factory=factory, start_worker=True)
    with TestClient(restarted) as client:
        assert client.get("/health/ready").status_code == 200
        assert client.get(f"/api/v1/documents/{document_ids[0]}").json()["status"] == "indexed"
        with _app(client).state.session_factory() as session:
            state = session.get_one(IndexState, 1)
            assert state.index_version == version_before_restart
        query_vector = encoder.encode([query_text], batch_size=1)
        scores, identifiers = _app(client).state.active_index.search(query_vector, 1)
        assert int(identifiers[0, 0]) == expected_vector_id
        assert float(scores[0, 0]) > 0.999
    assert factory.calls == 2


def test_parse_and_snapshot_failures_leave_previous_index_active(tmp_path: Path) -> None:
    settings = phase3_settings(tmp_path)
    encoder = DeterministicEncoder(settings.embedding_dimension)
    factory = CountingEncoderFactory(encoder)
    app = create_app(settings, encoder_factory=factory, start_worker=True)
    with TestClient(app) as client:
        seed_job, _ = _upload(
            client,
            "seed.txt",
            b"Seed evidence establishes a known good active snapshot.",
            "text/plain",
        )
        assert _wait_for_job(client, seed_job)["status"] == "succeeded"
        with _app(client).state.session_factory() as session:
            state = session.get_one(IndexState, 1)
            original_version = state.index_version
            original_count = state.vector_count

        parse_job, parse_document_id = _upload(
            client, "broken.pdf", b"%PDF-this-is-not-a-real-pdf", "application/pdf"
        )
        parse_failure = _wait_for_job(client, parse_job)
        assert parse_failure["status"] == "failed"
        assert parse_failure["error"] == {
            "code": "INVALID_FILE_CONTENT",
            "message": "The PDF could not be read safely.",
        }
        assert parse_failure["attempt"] == 1

        with patch(
            "app.services.index_coordinator.persist_snapshot",
            side_effect=OSError("simulated snapshot write failure"),
        ):
            index_job, index_document_id = _upload(
                client,
                "index-failure.txt",
                b"Distinct content reaches the atomic snapshot writer and then fails.",
                "text/plain",
            )
            index_failure = _wait_for_job(client, index_job)
        assert index_failure["status"] == "failed"
        index_error = cast(dict[str, object], index_failure["error"])
        assert index_error["code"] == "INDEX_UPDATE_FAILED"
        assert index_failure["attempt"] == 3

        failed_report = type(
            "FailedReport",
            (),
            {"ready": False, "errors": ("simulated post-commit mismatch",)},
        )()
        with patch(
            "app.services.index_coordinator.verify_active_snapshot",
            return_value=failed_report,
        ):
            verification_job, verification_document_id = _upload(
                client,
                "verification-failure.txt",
                b"A candidate snapshot is compensated when final alignment verification fails.",
                "text/plain",
            )
            verification_failure = _wait_for_job(client, verification_job)
        assert verification_failure["status"] == "failed"
        verification_error = cast(dict[str, object], verification_failure["error"])
        assert verification_error["code"] == "INDEX_UPDATE_FAILED"
        assert verification_failure["attempt"] == 3

        with _app(client).state.session_factory() as session:
            state = session.get_one(IndexState, 1)
            assert state.index_version == original_version
            assert state.vector_count == original_count
            assert session.scalar(select(func.count()).select_from(Chunk)) == original_count
            assert (
                session.scalar(
                    select(func.count())
                    .select_from(DocumentPage)
                    .where(
                        DocumentPage.document_id.in_(
                            [parse_document_id, index_document_id, verification_document_id]
                        )
                    )
                )
                == 0
            )
            failed_documents = session.scalars(
                select(Document).where(
                    Document.id.in_(
                        [parse_document_id, index_document_id, verification_document_id]
                    )
                )
            ).all()
            assert {document.status for document in failed_documents} == {"failed"}
        snapshots = [
            path for path in (settings.storage_root / "indexes").iterdir() if path.is_dir()
        ]
        assert len(snapshots) == 1
        assert client.get("/health/ready").status_code == 200


def test_restart_reconciles_and_retries_a_stale_running_job(tmp_path: Path) -> None:
    settings = phase3_settings(tmp_path)
    upgrade_database(settings.database_url)
    settings.storage_root.mkdir(parents=True, exist_ok=True)
    originals = settings.storage_root / "originals"
    originals.mkdir(parents=True, exist_ok=True)
    document_id = "d2097928-3ec4-4a0f-821c-31bc4b804587"
    job_id = "6f86e08e-7b97-493f-a865-e1c18a3a8fac"
    safe_name = f"{document_id}.txt"
    content = b"A restarted durable worker must retry this interrupted ingestion job."
    (originals / safe_name).write_bytes(content)
    engine = create_database_engine(settings.database_url)
    session_factory = create_session_factory(engine)
    with session_factory.begin() as session:
        session.add(
            Document(
                id=document_id,
                original_file_name="restarted.txt",
                safe_storage_name=safe_name,
                file_type="txt",
                mime_type="text/plain",
                domain="restart-test",
                source_kind="upload",
                storage_path=f"originals/{safe_name}",
                size_bytes=len(content),
                sha256=hashlib.sha256(content).hexdigest(),
                page_count=0,
                chunk_count=0,
                status="processing",
            )
        )
        session.add(
            IngestionJob(
                id=job_id,
                document_id=document_id,
                kind="ingest",
                status="running",
                stage="embedding",
                progress_percent=70,
                attempt=1,
                max_attempts=3,
            )
        )
    engine.dispose()

    encoder = DeterministicEncoder(settings.embedding_dimension)
    app = create_app(
        settings,
        encoder_factory=CountingEncoderFactory(encoder),
        start_worker=True,
    )
    with TestClient(app) as client:
        terminal = _wait_for_job(client, job_id)
        assert terminal["status"] == "succeeded"
        assert terminal["attempt"] == 2
        assert client.get(f"/api/v1/documents/{document_id}").json()["status"] == "indexed"
        assert client.get(f"/api/v1/documents/{document_id}/chunks").json()["total"] >= 1


def test_upload_contract_errors_are_typed(tmp_path: Path) -> None:
    settings = phase3_settings(tmp_path)
    disabled = create_app(settings, start_worker=False)
    with TestClient(disabled) as client:
        unavailable = client.post(
            "/api/v1/documents/upload",
            files={"file": ("sample.txt", b"content", "text/plain")},
        )
        assert unavailable.status_code == 503
        assert unavailable.json()["error"]["code"] == "WORKER_NOT_READY"
        missing_job = client.get("/api/v1/ingestion-jobs/00000000-0000-0000-0000-000000000000")
        assert missing_job.status_code == 404
        assert missing_job.json()["error"]["code"] == "JOB_NOT_FOUND"

    enabled = create_app(
        settings,
        encoder_factory=CountingEncoderFactory(DeterministicEncoder(4)),
        start_worker=True,
    )
    with TestClient(enabled) as client:
        unsupported = client.post(
            "/api/v1/documents/upload",
            files={"file": ("sample.exe", b"content", "application/octet-stream")},
        )
        assert unsupported.status_code == 415
        assert unsupported.json()["error"]["code"] == "UNSUPPORTED_FILE_TYPE"
        malformed = client.post(
            "/api/v1/documents/upload",
            files={"file": ("sample.pdf", b"plain text", "application/pdf")},
        )
        assert malformed.status_code == 422
        assert malformed.json()["error"]["code"] == "INVALID_FILE_CONTENT"
        invalid_metadata = client.post(
            "/api/v1/documents/upload",
            files={"file": ("metadata.txt", b"valid unique content", "text/plain")},
            data={"sourceUrl": "ftp://not-allowed.example/file"},
        )
        assert invalid_metadata.status_code == 422
        assert invalid_metadata.json()["error"]["code"] == "VALIDATION_ERROR"
        assert not list((settings.storage_root / "temp").glob("*.upload"))
    limited_settings = phase3_settings(tmp_path / "limited").model_copy(update={"max_upload_mb": 1})
    limited = create_app(
        limited_settings,
        encoder_factory=CountingEncoderFactory(DeterministicEncoder(4)),
        start_worker=True,
    )
    with TestClient(limited) as client:
        too_large = client.post(
            "/api/v1/documents/upload",
            files={"file": ("large.txt", b"x" * (1024 * 1024 + 1), "text/plain")},
        )
        assert too_large.status_code == 413
        assert too_large.json()["error"]["code"] == "FILE_TOO_LARGE"


def test_worker_model_load_failure_is_truthful_and_releases_queue_lock(tmp_path: Path) -> None:
    settings = phase3_settings(tmp_path)

    def fail_to_load(_settings: Settings) -> DeterministicEncoder:
        raise RuntimeError("simulated model load failure")

    with TestClient(
        create_app(settings, encoder_factory=fail_to_load, start_worker=True)
    ) as client:
        readiness = client.get("/health/ready").json()
        assert readiness["checks"]["embedding"]["ready"] is False
        assert readiness["checks"]["worker"]["ready"] is False
        response = client.post(
            "/api/v1/documents/upload",
            files={"file": ("sample.txt", b"content", "text/plain")},
        )
        assert response.status_code == 503
        assert response.json()["error"]["code"] == "WORKER_NOT_READY"
    lock = FileLock(settings.storage_root / "worker.lock")
    lock.acquire(timeout=0)
    lock.release()


def test_phase3_openapi_operations_match_contract(tmp_path: Path) -> None:
    settings = phase3_settings(tmp_path)
    with TestClient(create_app(settings, start_worker=False)) as client:
        openapi = client.get("/openapi.json").json()
    assert openapi["paths"]["/api/v1/documents/upload"]["post"]["operationId"] == ("uploadDocument")
    assert openapi["paths"]["/api/v1/ingestion-jobs/{jobId}"]["get"]["operationId"] == (
        "getIngestionJob"
    )
